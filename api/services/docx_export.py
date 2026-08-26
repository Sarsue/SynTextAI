"""A draft, as a Word document somebody can open.

Markdown to .docx, handling the shapes the drafting prompt actually asks the
model to produce: one H1 title, H2/H3 sections, ordered and unordered lists,
pipe tables, and bold/italic/code inline. Anything it does not recognise is
emitted as an ordinary paragraph rather than dropped, because losing a line of
a document a customer is about to hand to staff is worse than styling it
plainly.

There is no markdown-to-docx library in this project's dependencies and adding
one to convert six constructs would be a poor trade. python-docx is already here
because DocxProcessor reads Word files on the way in.

WHY THE PROVENANCE NOTE IS NOT OPTIONAL

A .docx leaves the product. It gets emailed, printed and pinned to a wall, and
that is exactly the moment everyone forgets a machine drafted it. The note is
written into the file itself, at the top, before the content.
"""
import re
from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Markdown that means something here. Everything else is a paragraph.
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_UNORDERED = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_ORDERED = re.compile(r"^(\s*)\d+[.)]\s+(.*)$")
_TABLE_SEPARATOR = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")
_RULE = re.compile(r"^\s*([-*_])\1{2,}\s*$")
# Bold before italic, so **x** is not read as *(*x*)*. Code last: its content is
# taken literally and must not be re-scanned for emphasis.
_INLINE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)")


def _add_runs(paragraph, text: str) -> None:
    """Write text into a paragraph, honouring bold, italic and inline code."""
    for part in _INLINE.split(text or ""):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**") and len(part) > 4) or \
           (part.startswith("__") and part.endswith("__") and len(part) > 4):
            paragraph.add_run(part[2:-2]).bold = True
        elif (part.startswith("*") and part.endswith("*") and len(part) > 2) or \
             (part.startswith("_") and part.endswith("_") and len(part) > 2):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def _split_row(line: str) -> List[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _write_table(document, rows: List[str]) -> None:
    """A pipe table, header row first.

    Falls back to plain paragraphs when the rows disagree about how many columns
    they have, which is what a half-written table looks like, rather than
    raising in the middle of somebody's download.
    """
    parsed = [_split_row(r) for r in rows]
    widths = {len(r) for r in parsed}
    if not parsed or len(widths) != 1 or parsed[0] == [""]:
        for r in rows:
            _add_runs(document.add_paragraph(), r)
        return

    table = document.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.style = "Table Grid"
    for i, row in enumerate(parsed):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_runs(paragraph, cell_text)
            if i == 0:
                for run in paragraph.runs:
                    run.bold = True


def _provenance(document, sources: Optional[List[Dict[str, Any]]]) -> None:
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run(
        "Written by SyntextAI from this workspace's documents, and edited by a "
        "person. Check it before anyone relies on it."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    names = []
    for s in (sources or []):
        name = s.get("file_name")
        if name and name not in names:
            names.append(name)
    if names:
        line = document.add_paragraph()
        run = line.add_run("Drawn from: " + ", ".join(names[:10]) + ".")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def markdown_to_docx(
    content: str,
    *,
    title: str = "",
    sources: Optional[List[Dict[str, Any]]] = None,
) -> bytes:
    """One draft as .docx bytes."""
    document = Document()

    if title:
        document.add_heading(title, level=0)
    _provenance(document, sources)

    lines = (content or "").split("\n")
    # The draft's title is normally taken from its own opening "# Heading", so
    # emitting that heading too would print the document's name twice, once as
    # Word's Title and again underneath it.
    for index, line in enumerate(lines):
        if not line.strip():
            continue
        if line.strip().startswith("# ") and line.strip()[2:].strip() == (title or "").strip():
            lines = lines[index + 1:]
        break
    table_rows: List[str] = []
    i = 0

    def flush_table():
        nonlocal table_rows
        if table_rows:
            _write_table(document, table_rows)
            table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # A table is the only construct that needs more than one line to
        # recognise: a pipe row is only a table once the next line is the
        # separator, otherwise it is prose that happens to contain a pipe.
        if stripped.startswith("|") and not table_rows:
            if i + 1 < len(lines) and _TABLE_SEPARATOR.match(lines[i + 1]):
                table_rows.append(stripped)
                i += 2  # header consumed, separator skipped
                continue
        elif table_rows:
            if stripped.startswith("|"):
                table_rows.append(stripped)
                i += 1
                continue
            flush_table()

        if not stripped:
            i += 1
            continue

        if _RULE.match(stripped):
            # A rule usually precedes the provenance footer the app appends, and
            # that is already at the top of this file. Emitting a second one
            # under a bare line looks like a mistake.
            i += 1
            continue

        heading = _HEADING.match(stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2).strip()
            # H1 is the document's own title. Word's Title style already stated
            # it above, so an H1 here becomes Heading 1 rather than a duplicate.
            document.add_heading(text, level=min(level, 4))
            i += 1
            continue

        ordered = _ORDERED.match(line)
        unordered = _UNORDERED.match(line)
        if ordered or unordered:
            match = ordered or unordered
            indent = len(match.group(1).expandtabs(4))
            style = "List Number" if ordered else "List Bullet"
            # Word names nested list styles "List Bullet 2", "List Bullet 3".
            depth = min(indent // 2, 2)
            if depth:
                style = f"{style} {depth + 1}"
            try:
                paragraph = document.add_paragraph(style=style)
            except KeyError:
                # A style the template does not carry. Plain paragraph rather
                # than a failed download.
                paragraph = document.add_paragraph()
            _add_runs(paragraph, match.group(2).strip())
            i += 1
            continue

        _add_runs(document.add_paragraph(), stripped)
        i += 1

    flush_table()

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def safe_filename(title: str, extension: str = "docx") -> str:
    """A filename a browser and a filesystem will both accept.

    The title is customer text and reaches a Content-Disposition header, where a
    quote or a newline would let it break out of the header value.
    """
    cleaned = re.sub(r"[^\w\s.-]", "", (title or "").strip(), flags=re.UNICODE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip().strip(".")
    if not cleaned:
        cleaned = "document"
    if len(cleaned) > 80:
        cleaned = cleaned[:80].rsplit(" ", 1)[0] or cleaned[:80]
    return f"{cleaned}.{extension}"
