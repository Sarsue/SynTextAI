"""A draft, as a file somebody can open: Word or PDF.

ONE PARSER, TWO RENDERERS

Markdown is read once into a list of blocks, and each format renders those
blocks its own way. The alternative was a second markdown parser inside the PDF
writer, and two parsers for one grammar drift: the day somebody teaches one of
them about nested lists, the other silently keeps flattening them, and nobody
finds out until a customer's SOP prints wrong.

WHAT IS SUPPORTED

The shapes the drafting prompt actually asks the model to produce: one H1 title,
H2/H3 sections, ordered and unordered lists with one level of nesting, pipe
tables, and bold/italic/code inline. Anything unrecognised becomes an ordinary
paragraph rather than being dropped, because losing a line of a document
somebody is about to hand to staff is worse than styling it plainly.

WHY THE PROVENANCE NOTE IS NOT OPTIONAL

Both formats leave the product. They get emailed, printed and pinned to a wall,
and that is exactly the moment everyone forgets a machine drafted it. The note
is written into the file itself, above the content, in both.
"""
import html as html_module
import re
from dataclasses import dataclass, field
from io import BytesIO
from typing import Any, Dict, List, Optional, Tuple

# Markdown that means something here. Everything else is a paragraph.
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_UNORDERED = re.compile(r"^(\s*)[-*+]\s+(.*)$")
_ORDERED = re.compile(r"^(\s*)\d+[.)]\s+(.*)$")
_TABLE_SEPARATOR = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")
_RULE = re.compile(r"^\s*([-*_])\1{2,}\s*$")
# Bold before italic, so **x** is not read as *(*x*)*. Code last: its content is
# taken literally and must not be re-scanned for emphasis.
_INLINE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_|`.+?`)")

PROVENANCE = (
    "Written by SyntextAI from this workspace's documents, and edited by a "
    "person. Check it before anyone relies on it."
)


@dataclass
class Block:
    kind: str                                   # heading | list | table | paragraph
    text: str = ""
    level: int = 0                              # heading only
    # (depth, ordered, text). Ordered is per ITEM, not per list, because an SOP
    # is normally a numbered procedure with unordered notes hanging off its
    # steps. Holding one flag for the whole block forced the parser to end the
    # list and start another at every switch, so the numbering restarted at 1
    # after every sub-item and a five-step procedure printed as 1, 1, 1, 1, 1.
    items: List[Tuple[int, bool, str]] = field(default_factory=list)
    rows: List[List[str]] = field(default_factory=list)          # table only


def _split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def parse_markdown(content: str) -> List[Block]:
    """Markdown into blocks. Never raises: a document must always export."""
    lines = (content or "").split("\n")
    blocks: List[Block] = []
    table_rows: List[str] = []
    list_items: List[Tuple[int, bool, str]] = []
    i = 0

    def flush_table():
        nonlocal table_rows
        if not table_rows:
            return
        parsed = [_split_row(r) for r in table_rows]
        widths = {len(r) for r in parsed}
        if len(widths) == 1 and parsed and parsed[0] != [""]:
            blocks.append(Block(kind="table", rows=parsed))
        else:
            # Rows disagreeing about their column count is what a truncated
            # table looks like. Plain text beats an exception in the middle of
            # somebody's download.
            for r in table_rows:
                blocks.append(Block(kind="paragraph", text=r))
        table_rows = []

    def flush_list():
        nonlocal list_items
        if list_items:
            blocks.append(Block(kind="list", items=list_items))
            list_items = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # A table needs two lines to recognise: a pipe row is only a table once
        # the next line is the separator, otherwise it is prose with a pipe in it.
        if stripped.startswith("|") and not table_rows:
            if i + 1 < len(lines) and _TABLE_SEPARATOR.match(lines[i + 1]):
                flush_list()
                table_rows.append(stripped)
                i += 2                      # header consumed, separator skipped
                continue
        elif table_rows:
            if stripped.startswith("|"):
                table_rows.append(stripped)
                i += 1
                continue
            flush_table()

        if not stripped:
            # A blank line does not necessarily end a list. Markdown allows a
            # loose list, where items are separated by blank lines, and ending
            # the list here would restart the numbering at the next item. Only
            # a blank line followed by something that is not a list item ends
            # one.
            if list_items:
                nxt = next((l for l in lines[i + 1:] if l.strip()), "")
                if not (_ORDERED.match(nxt) or _UNORDERED.match(nxt)):
                    flush_list()
            i += 1
            continue

        if _RULE.match(stripped):
            # A rule usually precedes the provenance footer the app appends to
            # an approved copy, and that note is already at the top of the file.
            flush_list()
            i += 1
            continue

        heading = _HEADING.match(stripped)
        if heading:
            flush_list()
            blocks.append(Block(
                kind="heading",
                level=min(len(heading.group(1)), 4),
                text=heading.group(2).strip(),
            ))
            i += 1
            continue

        ordered = _ORDERED.match(line)
        unordered = _UNORDERED.match(line)
        if ordered or unordered:
            match = ordered or unordered
            # No flush on a switch between numbered and bulleted. Numbered steps
            # with bulleted notes underneath are one list, and splitting them
            # is what made every step number 1.
            depth = min(len(match.group(1).expandtabs(4)) // 2, 2)
            list_items.append((depth, bool(ordered), match.group(2).strip()))
            i += 1
            continue

        flush_list()
        blocks.append(Block(kind="paragraph", text=stripped))
        i += 1

    flush_list()
    flush_table()
    return blocks


def drop_duplicate_title(blocks: List[Block], title: str) -> List[Block]:
    """Remove an opening H1 that just repeats the document's name.

    A draft normally takes its title from its own first heading, so rendering
    both prints the name twice: once as the file's title and again directly
    underneath it.
    """
    wanted = (title or "").strip()
    if not wanted or not blocks:
        return blocks
    first = blocks[0]
    if first.kind == "heading" and first.level == 1 and first.text.strip() == wanted:
        return blocks[1:]
    return blocks


def source_names(sources: Optional[List[Dict[str, Any]]], limit: int = 10) -> List[str]:
    """One name per document, not one per retrieved page.

    A twelve-page manual contributing eight passages is one source to a reader.
    """
    names: List[str] = []
    for s in (sources or []):
        name = s.get("file_name")
        if name and name not in names:
            names.append(name)
    return names[:limit]


# --- Word ---------------------------------------------------------------------

def _add_runs(paragraph, text: str) -> None:
    """Write text into a docx paragraph, honouring bold, italic and inline code."""
    for part in _INLINE.split(text or ""):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**") and len(part) > 4) or \
           (part.startswith("__") and part.endswith("__") and len(part) > 4):
            paragraph.add_run(part[2:-2]).bold = True
        elif (part.startswith("*") and part.endswith("*") and len(part) > 2) or \
             (part.startswith("_") and part.endswith("_") and len(part) > 2):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def markdown_to_docx(
    content: str,
    *,
    title: str = "",
    sources: Optional[List[Dict[str, Any]]] = None,
) -> bytes:
    """One draft as .docx bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    document = Document()

    if title:
        document.add_heading(title, level=0)

    def _note(text: str) -> None:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _note(PROVENANCE)
    names = source_names(sources)
    if names:
        _note("Drawn from: " + ", ".join(names) + ".")

    for block in drop_duplicate_title(parse_markdown(content), title):
        if block.kind == "heading":
            document.add_heading(block.text, level=block.level)
        elif block.kind == "list":
            for depth, ordered, text in block.items:
                style = "List Number" if ordered else "List Bullet"
                if depth:
                    style = f"{style} {depth + 1}"
                try:
                    paragraph = document.add_paragraph(style=style)
                except KeyError:
                    # A style this template does not carry. Plain paragraph
                    # rather than a failed download.
                    paragraph = document.add_paragraph()
                _add_runs(paragraph, text)
        elif block.kind == "table":
            table = document.add_table(rows=len(block.rows), cols=len(block.rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(block.rows):
                for c, cell_text in enumerate(row):
                    cell = table.cell(r, c)
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    _add_runs(paragraph, cell_text)
                    if r == 0:
                        for run in paragraph.runs:
                            run.bold = True
        else:
            _add_runs(document.add_paragraph(), block.text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- PDF ----------------------------------------------------------------------

def _inline_html(text: str) -> str:
    """Inline markdown as HTML, with the text escaped first.

    Escaping has to happen before the tags are added, or a document containing
    "a < b" would produce broken markup and MuPDF would render the rest of the
    paragraph as an unknown tag.
    """
    out = []
    for part in _INLINE.split(text or ""):
        if not part:
            continue
        if (part.startswith("**") and part.endswith("**") and len(part) > 4) or \
           (part.startswith("__") and part.endswith("__") and len(part) > 4):
            out.append(f"<b>{html_module.escape(part[2:-2])}</b>")
        elif (part.startswith("*") and part.endswith("*") and len(part) > 2) or \
             (part.startswith("_") and part.endswith("_") and len(part) > 2):
            out.append(f"<i>{html_module.escape(part[1:-1])}</i>")
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            out.append(f"<code>{html_module.escape(part[1:-1])}</code>")
        else:
            out.append(html_module.escape(part))
    return "".join(out)


def blocks_to_html(
    blocks: List[Block],
    *,
    title: str = "",
    sources: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """The document as the small HTML subset MuPDF's Story lays out."""
    parts: List[str] = ["<html><body>"]
    if title:
        parts.append(f"<h1>{html_module.escape(title)}</h1>")
    parts.append(f'<p class="note"><i>{html_module.escape(PROVENANCE)}</i></p>')
    names = source_names(sources)
    if names:
        joined = html_module.escape(", ".join(names))
        parts.append(f'<p class="note"><i>Drawn from: {joined}.</i></p>')

    for block in blocks:
        if block.kind == "heading":
            # The title above is the h1, so a heading inside the body starts at
            # h2 and never competes with it.
            tag = f"h{min(block.level + 1, 6)}"
            parts.append(f"<{tag}>{_inline_html(block.text)}</{tag}>")
        elif block.kind == "list":
            # A stack of the tags currently open, so a bulleted note nested
            # under a numbered step closes back into that step's <ol> and the
            # numbering carries on. Story has no indent property for list
            # items; nesting is expressed the way HTML expresses it, as a list
            # inside a list.
            open_tags: List[str] = []
            for depth, ordered, text in block.items:
                tag = "ol" if ordered else "ul"
                while len(open_tags) > depth + 1:
                    parts.append(f"</{open_tags.pop()}>")
                if len(open_tags) == depth + 1 and open_tags[-1] != tag:
                    # Same level, different kind: close this one and open the
                    # other, which is what markdown means there.
                    parts.append(f"</{open_tags.pop()}>")
                while len(open_tags) < depth + 1:
                    parts.append(f"<{tag}>")
                    open_tags.append(tag)
                parts.append(f"<li>{_inline_html(text)}</li>")
            while open_tags:
                parts.append(f"</{open_tags.pop()}>")
        elif block.kind == "table":
            parts.append("<table>")
            for r, row in enumerate(block.rows):
                cell_tag = "th" if r == 0 else "td"
                cells = "".join(f"<{cell_tag}>{_inline_html(c)}</{cell_tag}>" for c in row)
                parts.append(f"<tr>{cells}</tr>")
            parts.append("</table>")
        else:
            parts.append(f"<p>{_inline_html(block.text)}</p>")

    parts.append("</body></html>")
    return "".join(parts)


_PDF_CSS = """
body { font-family: sans-serif; font-size: 10.5pt; line-height: 1.4; }
h1 { font-size: 19pt; margin-bottom: 4pt; }
h2 { font-size: 14pt; margin-top: 12pt; margin-bottom: 3pt; }
h3 { font-size: 12pt; margin-top: 10pt; margin-bottom: 2pt; }
h4 { font-size: 11pt; margin-top: 8pt; }
p { margin-top: 0pt; margin-bottom: 6pt; }
p.note { font-size: 8.5pt; color: #666666; margin-bottom: 2pt; }
li { margin-bottom: 3pt; }
code { font-family: monospace; }
table { border: 1px solid #999999; }
th, td { border: 1px solid #999999; padding: 3pt; text-align: left; }
th { font-weight: bold; }
"""


def markdown_to_pdf(
    content: str,
    *,
    title: str = "",
    sources: Optional[List[Dict[str, Any]]] = None,
) -> bytes:
    """One draft as PDF bytes.

    MuPDF's Story lays out the same small HTML subset and paginates it, so no
    new dependency is needed: PyMuPDF is already here because PDFProcessor reads
    PDFs on the way in. A page is US Letter with one inch margins, which is what
    the customer's printer expects.

    The loop is Story's own contract: place() reports whether content is left
    over, and each pass draws what fitted onto a fresh page. It is capped
    because a block too tall to fit anywhere would otherwise never report
    finished, and an endless loop in a download is worse than a truncated
    document.
    """
    import fitz

    blocks = drop_duplicate_title(parse_markdown(content), title)
    html = blocks_to_html(blocks, title=title, sources=sources)

    story = fitz.Story(html=html, user_css=_PDF_CSS)
    buffer = BytesIO()
    writer = fitz.DocumentWriter(buffer)

    page = fitz.paper_rect("letter")
    frame = page + (72, 72, -72, -72)

    more = True
    pages = 0
    while more and pages < 200:
        device = writer.begin_page(page)
        more, _ = story.place(frame)
        story.draw(device)
        writer.end_page()
        pages += 1
    writer.close()
    return buffer.getvalue()


def safe_filename(title: str, extension: str = "docx") -> str:
    """A filename a browser and a filesystem will both accept.

    The title is customer text and reaches a Content-Disposition header, where a
    quote or a newline would let it break out of the header value.
    """
    cleaned = re.sub(r"[^\w\s.-]", "", (title or "").strip(), flags=re.UNICODE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip().strip(".")
    if not cleaned:
        cleaned = "document"
    if len(cleaned) > 80:
        cleaned = cleaned[:80].rsplit(" ", 1)[0] or cleaned[:80]
    return f"{cleaned}.{extension}"
