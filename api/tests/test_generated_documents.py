"""Documents SyntextAI writes, and the wall between them and the corpus.

The wall is the feature. If a generated draft could be retrieved before a person
approved it, the model's own output would become the model's own source of
truth: it writes a plausible SOP with one wrong figure, that gets ingested, and
afterwards it cites itself with a page reference indistinguishable from a real
one. The customer has no way to tell.

So the first tests here are not about generating anything. They are about a
draft being unable to answer a question, and about approval being the only door.
"""
import uuid

import pytest
import pytest_asyncio

from api.routes import drafts as drafts_route

pytestmark = pytest.mark.asyncio(loop_scope="session")

DIM = 1024
QUERY_VEC = [0.05] * DIM

SOURCE_TEXT = (
    "Sterilisation procedure. Instruments are cleaned in the ultrasonic bath "
    "for 10 minutes, then autoclaved at 134 degrees for 3 minutes."
)
# A figure that appears in no document, so retrieving it proves the draft was
# the source rather than the workspace.
DRAFT_TEXT = "# Sterilisation SOP\n\nAutoclave at 999 degrees for 42 minutes."


@pytest_asyncio.fixture(loop_scope="session")
async def paying(store, tenant):
    """Writing a document into a workspace is adding to it, so it is gated on
    the subscription exactly like an upload. Every test that generates or
    approves needs this."""
    from api.core.plans import STARTER

    await store.user_repo.add_or_update_subscription(
        user_id=tenant.owner,
        organization_id=tenant.org,
        stripe_customer_id="cus_test_drafts",
        stripe_subscription_id="sub_test_drafts",
        status="active",
        seats=STARTER.included_seats,
        plan_key="starter",
    )
    return tenant


@pytest_asyncio.fixture(loop_scope="session")
async def workspace(store, tenant):
    from api.models.orm_models import Chunk, Segment

    ws = await tenant.workspace("Drafting")
    file_id = await store.file_repo.add_file(
        user_id=tenant.owner, file_name=f"sop-{uuid.uuid4().hex[:6]}.pdf",
        file_url="", workspace_id=ws,
    )
    async with store.file_repo.get_async_session() as session:
        seg = Segment(file_id=file_id, page_number=1, content=SOURCE_TEXT)
        session.add(seg)
        await session.commit()
        session.add(Chunk(
            file_id=file_id, segment_id=seg.id, content=SOURCE_TEXT,
            embedding=QUERY_VEC, content_hash="h-" + uuid.uuid4().hex[:8],
        ))
        await session.commit()
    return {"id": ws, "file_id": file_id}


@pytest_asyncio.fixture(loop_scope="session")
async def draft(store, tenant, workspace):
    return await store.draft_repo.create(
        workspace_id=workspace["id"],
        created_by=tenant.owner,
        title="Sterilisation SOP",
        prompt="write a sterilisation SOP",
        content=DRAFT_TEXT,
        sources=[{"segment": 1, "file_id": workspace["file_id"],
                  "file_name": "sop.pdf", "page_number": 1}],
    )


# --- the wall ----------------------------------------------------------------

async def test_a_draft_cannot_answer_a_question(store, tenant, workspace, draft):
    """The whole point. A draft exists and retrieval cannot see it.

    Searched for a figure that appears only in the draft. If this ever returns
    a hit, a generated document has become its own source and the feature is
    actively harmful rather than merely incomplete.
    """
    hits = await store.file_repo.hybrid_search(
        user_id=tenant.owner,
        query="autoclave at 999 degrees for 42 minutes",
        query_embedding=QUERY_VEC,
        workspace_id=workspace["id"],
    )
    for h in hits:
        assert "999" not in (h.get("content") or ""), (
            "a generated draft was retrieved before anybody approved it"
        )


async def test_the_draft_is_not_a_file(store, tenant, workspace, draft):
    """It must not appear as a document either, or somebody will cite it by hand."""
    listed = await store.file_repo.get_files_for_user(
        tenant.owner, workspace_id=workspace["id"], limit=100
    )
    names = {f["file_name"] for f in listed["items"]}
    assert "Sterilisation SOP.md" not in names


# --- editing, which is why this is a document and not a chat message ---------

async def test_the_owner_can_edit_a_draft(store, tenant, draft, client):
    res = await client.as_(tenant.owner).patch(
        f"/api/v1/drafts/{draft['id']}",
        json={"content": "# Sterilisation SOP\n\nEdited by a person."},
    )
    assert res.status_code == 200, res.text
    assert "Edited by a person" in res.json()["content"]


async def test_an_empty_edit_is_refused(store, tenant, draft, client):
    res = await client.as_(tenant.owner).patch(f"/api/v1/drafts/{draft['id']}", json={})
    assert res.status_code == 400


async def test_the_list_does_not_carry_every_document_body(store, tenant, workspace, draft, client):
    """Twenty whole documents to render a list of titles is the mistake
    message_feedback was split out of `messages` to avoid."""
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts?workspace_id={workspace['id']}")
    assert res.status_code == 200, res.text
    items = res.json()["items"]
    assert items and "content" not in items[0]
    assert items[0]["title"] == "Sterilisation SOP"


# --- approval is the only door ------------------------------------------------

async def test_approving_creates_a_real_document_and_queues_ingestion(
    store, tenant, paying, workspace, draft, client, monkeypatch
):
    """Approval goes through the upload path rather than around it."""
    async def fake_upload(data, workspace_id, file_id, filename):
        fake_upload.seen = {"data": data, "filename": filename}
        return f"https://storage.example/{workspace_id}/{file_id}-{filename}"
    monkeypatch.setattr(drafts_route, "upload_bytes_to_gcs", fake_upload)

    res = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert res.status_code == 202, res.text
    file_id = res.json()["file_id"]

    created = await store.file_repo.get_file_by_id(file_id)
    assert created is not None
    assert created["workspace_id"] == workspace["id"]

    # The stored bytes say a machine wrote it and a person approved it. Once
    # this is retrievable it will be read by people who were not in the room.
    body = fake_upload.seen["data"].decode()
    assert "Written by SyntextAI" in body
    assert "approved by a person" in body

    refreshed = await store.draft_repo.get(draft["id"])
    assert refreshed["status"] == "ingested"
    assert refreshed["ingested_file_id"] == file_id


async def test_approving_twice_is_refused(
    store, tenant, paying, draft, client, monkeypatch
):
    async def fake_upload(data, workspace_id, file_id, filename):
        return f"https://storage.example/{file_id}-{filename}"
    monkeypatch.setattr(drafts_route, "upload_bytes_to_gcs", fake_upload)

    first = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert first.status_code == 202
    second = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert second.status_code == 409


async def test_a_failed_upload_leaves_no_unopenable_document(
    store, tenant, paying, workspace, draft, client, monkeypatch
):
    """A files row with no stored object is a document that appears in the list
    and can never be opened. Same rule as an upload."""
    async def failing_upload(data, workspace_id, file_id, filename):
        return None
    monkeypatch.setattr(drafts_route, "upload_bytes_to_gcs", failing_upload)

    before = await store.file_repo.get_files_for_user(
        tenant.owner, workspace_id=workspace["id"], limit=100
    )
    res = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert res.status_code == 500
    after = await store.file_repo.get_files_for_user(
        tenant.owner, workspace_id=workspace["id"], limit=100
    )
    assert after["total"] == before["total"]

    still_draft = await store.draft_repo.get(draft["id"])
    assert still_draft["status"] == "draft"


async def test_deleting_a_draft_keeps_the_document_approved_from_it(
    store, tenant, paying, draft, client, monkeypatch
):
    async def fake_upload(data, workspace_id, file_id, filename):
        return f"https://storage.example/{file_id}-{filename}"
    monkeypatch.setattr(drafts_route, "upload_bytes_to_gcs", fake_upload)

    ingested = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    file_id = ingested.json()["file_id"]

    deleted = await client.as_(tenant.owner).delete(f"/api/v1/drafts/{draft['id']}")
    assert deleted.status_code == 200

    assert await store.draft_repo.get(draft["id"]) is None
    assert await store.file_repo.get_file_by_id(file_id) is not None


# --- who may do it ------------------------------------------------------------

async def test_staff_cannot_edit_a_draft(store, tenant, workspace, draft, client):
    """Owners manage documents; staff ask questions."""
    member = await tenant.member(scope="workspace", workspaces=[workspace["id"]])
    res = await client.as_(member).patch(
        f"/api/v1/drafts/{draft['id']}", json={"content": "changed"},
    )
    assert res.status_code == 403


async def test_staff_cannot_approve_a_draft_into_the_knowledge_base(
    store, tenant, workspace, draft, client
):
    """The approval gate is the whole safety story, so it is held to the
    capability that adds documents, not to being able to read them."""
    member = await tenant.member(scope="workspace", workspaces=[workspace["id"]])
    res = await client.as_(member).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert res.status_code == 403


async def test_an_outsider_cannot_read_a_draft(store, tenant, draft, client):
    outsider = await tenant.new_user("outsider")
    res = await client.as_(outsider).get(f"/api/v1/drafts/{draft['id']}")
    assert res.status_code == 403


async def test_a_draft_that_does_not_exist_is_404(store, tenant, client):
    res = await client.as_(tenant.owner).get("/api/v1/drafts/99999999")
    assert res.status_code == 404


async def test_listing_drafts_in_an_unreachable_workspace_is_refused(
    store, tenant, workspace, client
):
    other = await tenant.workspace("Elsewhere")
    member = await tenant.member(scope="workspace", workspaces=[workspace["id"]])
    res = await client.as_(member).get(f"/api/v1/drafts?workspace_id={other}")
    assert res.status_code == 403


# --- generating ---------------------------------------------------------------

async def test_generating_refuses_when_no_document_covers_the_request(
    store, tenant, paying, workspace, client, monkeypatch
):
    """Writing a document from nothing is the failure this product exists
    against: fluent, confident, and sourced from the model's training data."""
    async def no_hits(**kwargs):
        return []
    monkeypatch.setattr(store.file_repo, "hybrid_search", no_hits)
    async def fake_embedding(text):
        return QUERY_VEC
    monkeypatch.setattr(drafts_route, "get_text_embedding", fake_embedding, raising=False)

    res = await client.as_(tenant.owner).post(
        "/api/v1/drafts/generate",
        json={"workspace_id": workspace["id"], "prompt": "write a fire safety policy"},
    )
    assert res.status_code == 422
    assert "no documents" in res.json()["detail"].lower()


async def test_a_generated_draft_is_saved_but_not_retrievable(
    store, tenant, paying, workspace, client, monkeypatch
):
    async def fake_embedding(text):
        return QUERY_VEC
    async def fake_generate(prompt, language=None, comprehension_level=None):
        return "# Fire Safety\n\nExtinguishers are checked every 77 months."
    monkeypatch.setattr(drafts_route, "get_text_embedding", fake_embedding, raising=False)
    monkeypatch.setattr(drafts_route, "generate_explanation", fake_generate)

    res = await client.as_(tenant.owner).post(
        "/api/v1/drafts/generate",
        json={"workspace_id": workspace["id"], "prompt": "write a sterilisation SOP"},
    )
    assert res.status_code == 201, res.text
    body = res.json()
    assert body["status"] == "draft"
    assert body["sources"], "a draft with no recorded sources cannot be checked"

    hits = await store.file_repo.hybrid_search(
        user_id=tenant.owner,
        query="extinguishers checked every 77 months",
        query_embedding=QUERY_VEC,
        workspace_id=workspace["id"],
    )
    for h in hits:
        assert "77 months" not in (h.get("content") or "")


async def test_staff_cannot_generate(store, tenant, workspace, client):
    member = await tenant.member(scope="workspace", workspaces=[workspace["id"]])
    res = await client.as_(member).post(
        "/api/v1/drafts/generate",
        json={"workspace_id": workspace["id"], "prompt": "write a sterilisation SOP"},
    )
    assert res.status_code == 403


def test_a_title_is_made_from_the_request_without_an_llm_call():
    assert drafts_route._title_from("write a sterilisation SOP") == "Sterilisation SOP"
    assert drafts_route._title_from("Create an onboarding checklist") == "Onboarding checklist"
    assert drafts_route._title_from("") == "Untitled document"


async def test_a_draft_can_be_approved_again_if_the_document_was_deleted(
    store, tenant, paying, draft, client, monkeypatch
):
    """Deleting the approved document must not strand the draft.

    ingested_file_id is ON DELETE SET NULL while status stays 'ingested', so
    reading status alone would refuse a second approval forever and leave a
    draft that says it is in the knowledge base while nothing is.
    """
    async def fake_upload(data, workspace_id, file_id, filename):
        return f"https://storage.example/{file_id}-{filename}"
    monkeypatch.setattr(drafts_route, "upload_bytes_to_gcs", fake_upload)

    first = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert first.status_code == 202
    file_id = first.json()["file_id"]

    await store.file_repo.delete_file_entry(file_id)
    orphaned = await store.draft_repo.get(draft["id"])
    assert orphaned["ingested_file_id"] is None, "the FK did not clear"

    again = await client.as_(tenant.owner).post(f"/api/v1/drafts/{draft['id']}/ingest")
    assert again.status_code == 202, again.text
    assert again.json()["file_id"] != file_id


# --- export -------------------------------------------------------------------

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _docx_text(payload: bytes):
    """Every paragraph and table cell in a .docx, so a test can read what a
    person would see rather than trusting the byte count."""
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(payload))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.extend(c.text for c in row.cells)
    return doc, lines


async def test_a_draft_downloads_as_a_word_document(store, tenant, draft, client):
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts/{draft['id']}/export")
    assert res.status_code == 200, res.text
    assert res.headers["content-type"] == DOCX_MEDIA_TYPE
    assert "attachment" in res.headers["content-disposition"]
    # A real Word file, opened rather than assumed.
    _, lines = _docx_text(res.content)
    assert any("Sterilisation SOP" in l for l in lines)
    assert any("999 degrees" in l for l in lines)


async def test_the_downloaded_file_says_a_machine_wrote_it(store, tenant, draft, client):
    """The note travels with the file.

    A .docx leaves the product: it gets emailed, printed and pinned to a wall,
    and that is exactly when everybody forgets a machine drafted it. Marking it
    only in the app would mark it in the one place it is already obvious.
    """
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts/{draft['id']}/export")
    _, lines = _docx_text(res.content)
    assert any("Written by SyntextAI" in l for l in lines)
    assert any("sop.pdf" in l for l in lines), "the source documents are not named"


async def test_headings_lists_and_tables_survive_the_conversion(
    store, tenant, workspace, client
):
    """Word structure, not one long paragraph. A document somebody has to
    reformat before using is a document they will rewrite instead."""
    from docx import Document as DocxDocument
    from io import BytesIO

    rich = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Rich", prompt="p",
        content=(
            "# Rich\n\n"
            "## Steps\n\n"
            "1. First step\n"
            "2. Second step\n\n"
            "- A bullet\n"
            "- Another bullet\n\n"
            "## Limits\n\n"
            "| Setting | Value |\n"
            "|---|---|\n"
            "| Temperature | 134 |\n"
            "| Time | 3 min |\n\n"
            "Some **bold** text and some *italic* text.\n"
        ),
        sources=[],
    )
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts/{rich['id']}/export")
    assert res.status_code == 200, res.text

    doc = DocxDocument(BytesIO(res.content))
    styles = [p.style.name for p in doc.paragraphs]
    assert any(s.startswith("Heading") for s in styles), "no headings survived"
    assert any(s.startswith("List Number") for s in styles), "numbered list lost"
    assert any(s.startswith("List Bullet") for s in styles), "bullet list lost"

    assert len(doc.tables) == 1, f"expected one table, got {len(doc.tables)}"
    cells = [c.text for r in doc.tables[0].rows for c in r.cells]
    assert "Temperature" in cells and "134" in cells

    # The markers themselves must not survive as literal text.
    body = "\n".join(p.text for p in doc.paragraphs)
    assert "**bold**" not in body and "bold" in body
    assert "|---|" not in body


async def test_a_half_written_table_becomes_text_rather_than_an_error(
    store, tenant, workspace, client
):
    """Rows disagreeing about their column count is what a truncated table looks
    like, and a download that raises is worse than one that is plain."""
    broken = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Broken", prompt="p",
        content="| A | B |\n|---|---|\n| only one\n",
        sources=[],
    )
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts/{broken['id']}/export")
    assert res.status_code == 200, res.text
    _, lines = _docx_text(res.content)
    assert any("only one" in l for l in lines), "the row was dropped"


async def test_staff_can_download_what_they_can_read(
    store, tenant, workspace, draft, client
):
    """Somebody who may see the document may take a copy. Refusing that while
    showing them the text on screen would be theatre."""
    member = await tenant.member(scope="workspace", workspaces=[workspace["id"]])
    res = await client.as_(member).get(f"/api/v1/drafts/{draft['id']}/export")
    assert res.status_code == 200


async def test_an_outsider_cannot_download_a_draft(store, tenant, draft, client):
    outsider = await tenant.new_user("outsider")
    res = await client.as_(outsider).get(f"/api/v1/drafts/{draft['id']}/export")
    assert res.status_code == 403


def test_a_title_cannot_break_out_of_the_content_disposition_header():
    """The title is customer text and reaches a header. A quote or a newline
    there would let it inject header fields."""
    from api.services.document_export import safe_filename

    nasty = safe_filename('evil"; drop\r\nSet-Cookie: a=b')
    assert '"' not in nasty and "\r" not in nasty and "\n" not in nasty
    assert nasty.endswith(".docx")
    assert safe_filename("") == "document.docx"
    assert safe_filename("   ...   ") == "document.docx"


async def test_the_title_comes_from_the_document_not_the_request(
    store, tenant, paying, workspace, client, monkeypatch
):
    """"A hand hygiene quick reference for new staff. Use a table for when..."
    is what somebody typed. "Hand Hygiene Quick Reference" is what they wanted,
    and it is also the filename of the .docx they download."""
    async def fake_embedding(text):
        return QUERY_VEC
    async def fake_generate(prompt, language=None, comprehension_level=None):
        return "# Hand Hygiene Quick Reference\n\n## When to wash\n\nWhen visibly soiled."
    monkeypatch.setattr(drafts_route, "get_text_embedding", fake_embedding, raising=False)
    monkeypatch.setattr(drafts_route, "generate_explanation", fake_generate)

    res = await client.as_(tenant.owner).post(
        "/api/v1/drafts/generate",
        json={"workspace_id": workspace["id"],
              "prompt": "A hand hygiene quick reference for new staff. Use a table "
                        "for when to wash versus when to use alcohol rub."},
    )
    assert res.status_code == 201, res.text
    assert res.json()["title"] == "Hand Hygiene Quick Reference"


async def test_the_document_name_is_not_printed_twice(
    store, tenant, workspace, client
):
    """Word's Title style already states it, so the opening # heading would be
    the same words again directly underneath."""
    from docx import Document as DocxDocument
    from io import BytesIO

    d = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Hand Hygiene", prompt="p",
        content="# Hand Hygiene\n\n## When to wash\n\nWhen visibly soiled.",
        sources=[],
    )
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts/{d['id']}/export")
    doc = DocxDocument(BytesIO(res.content))
    exact = [p.text for p in doc.paragraphs if p.text.strip() == "Hand Hygiene"]
    assert len(exact) == 1, f"the title appears {len(exact)} times"


def test_the_prompt_forbids_citing_segment_numbers_in_the_output():
    """The model wrote "(Segment 6)" into a procedure on the first real run.

    Segment numbering is scaffolding for the model. A printed SOP that cites one
    reads as broken to the person following it, who has never heard of a
    segment.
    """
    built = drafts_route._draft_prompt("write an SOP", "[Segment 1]\nsome text", "")
    assert "Segment 4" in built or "Segment 9" in built, (
        "the instruction no longer shows the model what not to write"
    )
    assert "Do not write" in built


# --- PDF ----------------------------------------------------------------------

def _pdf_text(payload: bytes):
    """Every page's text, read back the way any PDF reader would."""
    import fitz

    doc = fitz.open(stream=payload, filetype="pdf")
    return doc, "\n".join(page.get_text() for page in doc)


async def test_a_draft_downloads_as_a_pdf(store, tenant, draft, client):
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{draft['id']}/export?format=pdf"
    )
    assert res.status_code == 200, res.text
    assert res.headers["content-type"] == "application/pdf"
    assert ".pdf" in res.headers["content-disposition"]
    # A real PDF, opened rather than assumed.
    assert res.content[:5] == b"%PDF-", "not a PDF at all"

    doc, text = _pdf_text(res.content)
    assert len(doc) >= 1
    assert "Sterilisation SOP" in text
    assert "999 degrees" in text


async def test_the_pdf_says_a_machine_wrote_it(store, tenant, draft, client):
    """Same reason as the Word file, more so: a PDF is the one people print."""
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{draft['id']}/export?format=pdf"
    )
    _, text = _pdf_text(res.content)
    assert "Written by SyntextAI" in text
    assert "sop.pdf" in text, "the source documents are not named"


async def test_headings_lists_and_tables_reach_the_pdf(store, tenant, workspace, client):
    rich = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Rich", prompt="p",
        content=(
            "# Rich\n\n## Steps\n\n1. First step\n2. Second step\n\n"
            "- A bullet\n\n## Limits\n\n"
            "| Setting | Value |\n|---|---|\n| Temperature | 134 |\n| Time | 3 min |\n\n"
            "Some **bold** text.\n"
        ),
        sources=[],
    )
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{rich['id']}/export?format=pdf"
    )
    assert res.status_code == 200, res.text
    _, text = _pdf_text(res.content)

    for expected in ("Steps", "First step", "A bullet", "Temperature", "134", "bold"):
        assert expected in text, f"{expected!r} did not reach the PDF"
    # The markers themselves must not survive as literal text.
    assert "**bold**" not in text
    assert "|---|" not in text


async def test_a_long_document_paginates(store, tenant, workspace, client):
    """Story lays out and breaks pages itself; this proves the loop that drives
    it actually asks for more than one."""
    long_doc = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Long", prompt="p",
        content="\n\n".join(
            f"## Section {n}\n\nParagraph {n}. " + ("Filler sentence. " * 40)
            for n in range(1, 25)
        ),
        sources=[],
    )
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{long_doc['id']}/export?format=pdf"
    )
    assert res.status_code == 200, res.text
    doc, text = _pdf_text(res.content)
    assert len(doc) > 1, "a very long document came out as one page"
    assert "Section 24" in text, "the end of the document was lost"


async def test_angle_brackets_in_the_text_do_not_break_the_pdf(
    store, tenant, workspace, client
):
    """Text is escaped before markdown becomes HTML.

    Escaping afterwards would turn "a < b" into markup MuPDF reads as an unknown
    tag, and the rest of the paragraph disappears. The drafting prompt asks for
    "TO BE COMPLETED: <what is missing>", so this is the normal case, not an
    exotic one.
    """
    tricky = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Tricky", prompt="p",
        content="TO BE COMPLETED: <the autoclave temperature>\n\nAlso 5 < 6 & 7 > 2.\n",
        sources=[],
    )
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{tricky['id']}/export?format=pdf"
    )
    assert res.status_code == 200, res.text
    _, text = _pdf_text(res.content)
    assert "the autoclave temperature" in text, "an angle-bracketed gap vanished"
    assert "Also 5 < 6 & 7 > 2." in text


async def test_an_unknown_format_is_refused(store, tenant, draft, client):
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{draft['id']}/export?format=rtf"
    )
    assert res.status_code == 400


async def test_the_default_format_is_still_word(store, tenant, draft, client):
    """The button that shipped first sends no format parameter."""
    res = await client.as_(tenant.owner).get(f"/api/v1/drafts/{draft['id']}/export")
    assert res.status_code == 200
    assert res.headers["content-type"] == DOCX_MEDIA_TYPE


async def test_an_outsider_cannot_download_a_pdf(store, tenant, draft, client):
    outsider = await tenant.new_user("outsider")
    res = await client.as_(outsider).get(
        f"/api/v1/drafts/{draft['id']}/export?format=pdf"
    )
    assert res.status_code == 403


def test_both_formats_read_the_same_markdown():
    """One parser, two renderers. A second parser inside the PDF writer would
    drift: teach one about nested lists and the other keeps flattening them."""
    from api.services import document_export

    blocks = document_export.parse_markdown(
        "# T\n\n## S\n\n1. one\n2. two\n\n| A | B |\n|---|---|\n| 1 | 2 |\n"
    )
    kinds = [b.kind for b in blocks]
    assert kinds == ["heading", "heading", "list", "table"], kinds
    assert len(blocks[2].items) == 2 and all(o for _, o, _ in blocks[2].items)
    assert blocks[3].rows == [["A", "B"], ["1", "2"]]


async def test_a_numbered_procedure_with_sub_bullets_keeps_counting(
    store, tenant, workspace, client
):
    """The bug this closes, seen in a real generated SOP.

    Numbered steps with bulleted notes under them are how every procedure is
    written. Holding one ordered flag per list forced the parser to end the list
    and start another at each switch, so a five-step procedure printed as
    1, 1, 1, 1, 1 and the person following it had no idea what order to work in.
    """
    from api.services import document_export

    markdown = (
        "1. **Gloves** - Wear gloves whenever there is contact with blood.\n"
        "   - Do not reuse gloves.\n"
        "   - Change gloves between patients.\n"
        "2. **Protective clothing** - Wear a gown that covers skin.\n"
        "   - Change it if visibly soiled.\n"
        "3. **Eye protection** - Wear it during procedures that splash.\n"
    )

    blocks = document_export.parse_markdown(markdown)
    lists = [b for b in blocks if b.kind == "list"]
    assert len(lists) == 1, f"one procedure became {len(lists)} lists"
    top = [t for d, o, t in lists[0].items if d == 0 and o]
    assert len(top) == 3, f"expected 3 numbered steps, got {len(top)}"

    # And the HTML the PDF is laid out from keeps them in one <ol>.
    html = document_export.blocks_to_html(blocks)
    assert html.count("<ol>") == 1, "the numbered steps were split across lists"
    assert html.count("<ul>") == 2, "the notes did not nest under their steps"

    d = await store.draft_repo.create(
        workspace_id=workspace["id"], created_by=tenant.owner,
        title="Procedure", prompt="p", content=markdown, sources=[],
    )
    res = await client.as_(tenant.owner).get(
        f"/api/v1/drafts/{d['id']}/export?format=pdf"
    )
    assert res.status_code == 200, res.text
    _, text = _pdf_text(res.content)
    for n in ("1.", "2.", "3."):
        assert n in text, f"step {n} is not numbered in the PDF"


async def test_a_loose_list_does_not_restart_at_one(store, tenant, workspace, client):
    """Markdown allows blank lines between items. Ending the list there would
    restart the numbering at the next one."""
    from api.services import document_export

    blocks = document_export.parse_markdown(
        "1. First step\n\n2. Second step\n\n3. Third step\n"
    )
    lists = [b for b in blocks if b.kind == "list"]
    assert len(lists) == 1, f"blank lines split one list into {len(lists)}"
    assert len(lists[0].items) == 3


def test_a_list_still_ends_when_the_document_moves_on():
    """The blank-line rule must not swallow everything after a list."""
    from api.services import document_export

    blocks = document_export.parse_markdown(
        "1. First step\n2. Second step\n\nA following paragraph.\n\n- A later bullet\n"
    )
    kinds = [b.kind for b in blocks]
    assert kinds == ["list", "paragraph", "list"], kinds
