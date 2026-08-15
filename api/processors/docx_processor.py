"""
DOCX processor module - Handles extraction and processing of Word documents.
Mirrors PDFProcessor's shape so chunks flow through the same downstream
storage/retrieval path. See pdf_processor.py for the reference pattern.
"""
import logging
import gc
from io import BytesIO
from typing import Dict, List, Any

from docx import Document as DocxDocument

from api.repositories.repository_manager import RepositoryManager
from api.processors.base_processor import FileProcessor
from api.services.llm_service import get_text_embeddings_in_batches
from api.core.utils import chunk_text

logger = logging.getLogger(__name__)

HEADING_STYLES = {"Heading 1", "Heading 2", "Title"}


def _cell(text: str) -> str:
    """One cell's text, safe to sit between two pipes.

    A cell holding its own pipe, or a line break, would otherwise invent columns
    and rows that are not in the document.
    """
    return " ".join((text or "").split()).replace("|", r"\|")


def _markdown_table(table) -> List[str]:
    """A Word table as markdown rows: header, separator, then the body.

    WHY THIS SHAPE AND NOT " | ".join(cells)

    The rows used to be joined with a bare " | ", with no pipes on the ends and
    no separator line. That is not a table to anything downstream. `chunk_text`
    counts to 400 tokens and cuts wherever it lands, which on a table is
    mid-row and away from the header, and the pieces after the first carry no
    column names at all. Measured on a reconstructed charging chart 2026-08-13,
    that turned one table into three chunks of which all three were
    unanswerable, and it is what `chunk_markdown` and `_chunk_table` were built
    to prevent. They just never saw a DOCX, because only PDF pages read by the
    vision model were marked as markdown.

    Word has no notion of a header row, so the first row stands in for one.
    That is what it almost always is, and being wrong about it costs a repeated
    first row rather than a lost one.

    Rows are padded to the widest row. A merged cell makes python-docx repeat
    the same cell across the row, so widths can differ, and a ragged markdown
    table stops being one.
    """
    rows: List[List[str]] = []
    for row in table.rows:
        cells = [_cell(c.text) for c in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return []

    width = max(len(r) for r in rows)
    lines = [
        "| " + " | ".join(r + [""] * (width - len(r))) + " |"
        for r in rows
    ]
    return [lines[0], "|" + " --- |" * width] + lines[1:]


class DocxProcessor(FileProcessor):
    """
    Processor for Word (.docx) documents.
    Handles text extraction and embedding generation.
    """

    def __init__(self, store: RepositoryManager):
        """
        Initialize the DOCX processor.

        Args:
            store: RepositoryManager instance for database operations
        """
        super().__init__()
        self.store = store

    async def process(self,
                     file_data: bytes,
                     file_id: int,
                     user_id: int,
                     filename: str,
                     **kwargs) -> Dict[str, Any]:
        """
        Process a DOCX file: extract text, generate embeddings.

        Args:
            file_data: Raw DOCX file data in bytes
            file_id: Database ID of the file
            user_id: ID of the user who owns the file
            filename: Name of the file
            **kwargs: Additional arguments

        Returns:
            Dictionary containing processing results
        """
        logger.info(f"Processing DOCX file: {filename} (ID: {file_id}, User: {user_id})")

        section_data = self.extract_text_with_sections(file_data)
        logger.info(f"DOCX extraction complete. Sections: {len(section_data)}")

        if not section_data:
            logger.error(f"Failed to extract content from DOCX: {filename}")
            return {
                "success": False,
                "file_id": file_id,
                "error": "Failed to extract content from DOCX",
                "metadata": {
                    "processor_type": "docx"
                }
            }

        logger.info(f"Starting section processing and embedding generation for file {filename}")
        try:
            await self.store.file_repo.update_file_status(int(file_id), "embedding")
        except Exception:
            logger.debug("Non-fatal: could not update status to 'embedding'")
        # Chunks are written batch by batch inside this call, so a crash keeps
        # the sections it already reached and the answer can already cite them.
        result = await self.embed_and_store_pages(
            section_data,
            file_id=int(file_id),
            user_id=user_id,
            filename=filename,
            file_type="docx",
        )

        # A document that extracted nothing is a failure, not a success. A
        # processor could catch every section, produce zero chunks, and still
        # mark the file processed, leaving a document that looks ready in the
        # list and answers nothing. Counted across the whole document, not this
        # attempt: a resumed run that finds everything already stored did no
        # work and is finished, not empty.
        if result["stored_chunks"] == 0 and result["skipped_pages"] == 0:
            logger.error(
                f"No chunks extracted from {filename}; marking it failed rather "
                f"than leaving a document that looks ready and answers nothing"
            )
            try:
                await self.store.file_repo.update_file_status(int(file_id), "failed")
            except Exception:
                logger.debug("Non-fatal: could not update status to 'failed'")
            return {
                "success": False,
                "file_id": file_id,
                "error": "No content could be extracted from this document",
                "metadata": {"processor_type": "docx"},
            }

        return {
            "success": True,
            "file_id": file_id,
            "metadata": {
                "section_count": len(section_data),
                "chunk_count": result["stored_chunks"],
                "resumed_pages": result["skipped_pages"],
                "processor_type": "docx"
            }
        }

    def extract_text_with_sections(self, docx_data: bytes) -> List[Dict[str, Any]]:
        """
        Extracts text from DOCX data, split into logical sections at heading
        paragraphs (Heading 1/2/Title styles). Word has no fixed page concept
        outside of rendering, so sections stand in for "page_num" the same
        way PDF page numbers anchor a citation, "Section 3" instead of "p.3".
        Falls back to a single section if no headings are found. Table
        content is appended as trailing sections since python-docx exposes
        paragraphs and tables as separate lists, not interleaved in reading
        order, precise position isn't preserved, but the content isn't lost.
        Each table is rendered as markdown and marked as such, so it is cut on
        row boundaries rather than at a token count. See _markdown_table.

        Args:
            docx_data: DOCX file data in bytes

        Returns:
            List of dictionaries with section numbers and text content
        """
        try:
            doc = DocxDocument(BytesIO(docx_data))
        except Exception as e:
            logger.error(f"Error opening DOCX (may be corrupt or not a valid .docx): {e}", exc_info=True)
            return []

        sections: List[Dict[str, Any]] = []
        current_heading = None
        current_lines: List[str] = []

        def flush_section():
            content = "\n".join(current_lines).strip()
            if not content and current_heading:
                # A heading with no body text before the next heading (e.g. a
                # title immediately followed by the first section) would
                # otherwise be silently dropped, fall back to the heading
                # text itself so nothing gets lost.
                content = current_heading
            if content:
                section_num = len(sections) + 1
                label = f"Section {section_num}"
                if current_heading:
                    label += f": {current_heading}"
                sections.append({
                    "page_num": section_num,
                    "text": f"{label}\n{content}"
                })

        try:
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if para.style is not None and para.style.name in HEADING_STYLES:
                    flush_section()
                    current_heading = text
                    current_lines = []
                else:
                    current_lines.append(text)
            flush_section()

            # Tables aren't interleaved with paragraphs in python-docx's flat
            # lists, append them as their own trailing sections rather than
            # dropping them.
            for t_idx, table in enumerate(doc.tables, 1):
                rows_text = _markdown_table(table)
                if rows_text:
                    section_num = len(sections) + 1
                    sections.append({
                        "page_num": section_num,
                        # "Table N" on its own line above the table is the
                        # caption, and _split_markdown_blocks carries the line
                        # above a table into every piece it is cut into. It is
                        # the only thing naming a chunk of bare numbers, since
                        # a trailing table section has lost the heading it sat
                        # under.
                        "text": f"Table {t_idx}\n" + "\n".join(rows_text),
                        # Cut on row boundaries, with the caption and the header
                        # repeated above every piece, instead of at 400 tokens
                        # wherever that falls.
                        "is_markdown": True,
                    })

            if not sections:
                logger.warning("DOCX had no extractable paragraph or table text")

            return sections
        except Exception as e:
            logger.error(f"Error extracting DOCX sections: {e}", exc_info=True)
            return []

    async def extract_content(self, **kwargs) -> Dict[str, Any]:
        """
        Extract raw content from the DOCX file.

        Args:
            **kwargs: Must include 'file_data' as bytes

        Returns:
            Dict containing extracted section content
        """
        file_data = kwargs.get('file_data')
        if not file_data:
            raise ValueError("Missing required 'file_data' parameter")

        section_data = self.extract_text_with_sections(file_data)
        return {"pages": section_data}

    async def generate_embeddings(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate embeddings for the extracted DOCX content.

        Args:
            content: Dictionary containing sections with extracted content

        Returns:
            Dict containing content with embeddings
        """
        pages = content.get("pages", [])
        processed_data = await self.process_pages(pages)
        return processed_data

    def _log_error(self, message: str, error: Exception) -> None:
        """Log an error with consistent format."""
        logging.error(f"{message}: {str(error)[:200]}", exc_info=True)
